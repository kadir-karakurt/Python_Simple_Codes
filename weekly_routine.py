from docx import Document

# Create a Word document for weeks tasks
doc = Document()
doc.add_heading('Weekly Routine', 0)

# Creating a table
num_cols = 8  # 8 columns: Time, Monday, Tuesday, Wednesday, Thursday, Friday, Saturday, Sunday
table = doc.add_table(rows=1, cols=num_cols)
hdr_cells = table.rows[0].cells

# Adding headers
headers = ['Time', 'Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday', 'Sunday']
for i, header in enumerate(headers):
    hdr_cells[i].text = header

# Adding empty rows for weekly tasks (e.g., 10 rows)
for _ in range(10):
    row_cells = table.add_row().cells
    for i in range(num_cols):
        row_cells[i].text = ''

# Saving the document
file_path = "weekly_routine.docx"
doc.save(file_path)


# pip install python-docx
# run - python weekly_routine.py
